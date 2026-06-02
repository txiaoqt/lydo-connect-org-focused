from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("LYDO_Connect_PWA_Panel_Explanation.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK_BLUE = RGBColor(26, 63, 122)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
SOFT_YELLOW = "FFF4CC"


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = BLUE
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = BLUE
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = DARK_BLUE
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(5)
    heading_3.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    number = styles["List Number"]
    number.font.name = "Calibri"
    number._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    number.font.size = Pt(11)
    number.paragraph_format.left_indent = Inches(0.375)
    number.paragraph_format.first_line_indent = Inches(-0.188)
    number.paragraph_format.space_after = Pt(4)
    number.paragraph_format.line_spacing = 1.25


def set_footer(document: Document) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("LYDO Connect - PWA Panel Explanation")
    set_font(run, size=9, italic=True, color=MUTED)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120, width_dxa=9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")

    set_cell_margins(table)


def format_table_text(table, header_rows=1) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=10.5, bold=row_idx < header_rows)
            if row_idx < header_rows:
                shade_cell(cell, LIGHT_BLUE)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("LYDO CONNECT PWA EXPLANATION")
    set_font(run, size=22, bold=True, color=INK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Simple Panelist Q&A Script and Defense Guide")
    set_font(run, size=12, italic=True, color=MUTED)

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    rows = [
        ("Purpose", "Explain why LYDO Connect is a Progressive Web App (PWA) in simple, panel-friendly language."),
        ("Main Point", "The site can still be used in the browser, but PWA support gives users the option to use it like an app."),
        ("Best Audience", "Panelists, teachers, admins, staff, and youth users who need a quick explanation."),
        ("Honest Caveat", "Installation is optional, and online data features still need internet access."),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], LIGHT_GRAY)
    format_table_text(table, header_rows=0)
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, size=10.5, bold=True, color=DARK_BLUE)


def add_callout(document: Document, label: str, text: str, fill=SOFT_YELLOW) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=11, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_font(text_run, size=11)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="List Bullet")
    for run in paragraph.runs:
        set_font(run, size=11)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="List Number")
    for run in paragraph.runs:
        set_font(run, size=11)


def add_script_block(document: Document, title: str, script: str) -> None:
    document.add_heading(title, level=3)
    add_callout(document, "Say this", script, fill="F4F6F9")


def add_simple_pwa_explanation(document: Document) -> None:
    document.add_heading("1. Simple Meaning of PWA", level=1)
    add_callout(
        document,
        "Core sentence",
        "A PWA is a website that can behave like a mobile app when the user wants that experience.",
    )
    document.add_paragraph(
        "Very simple explanation: LYDO Connect is still a website, but it has app-like capabilities. "
        "A user can open it in the browser, and supported browsers can also let the user add it to the home screen or install it as an app shortcut."
    )
    document.add_paragraph(
        "ELI5 version: PWA is like giving the website an app costume. Website pa rin siya, pero kapag ginamit ng user, feeling niya app ang gamit niya."
    )

    document.add_heading("What the PWA adds", level=2)
    for item in [
        "It can have its own icon on the phone or computer home screen.",
        "It can open in a standalone app-like window instead of looking like a normal browser tab.",
        "It can load faster because important files can be cached by the service worker.",
        "It can still show the app shell or cached pages during weak or limited internet connection.",
        "It updates through the web when the site is updated, so users do not need Play Store or App Store updates.",
        "It works across phones, tablets, laptops, and desktops through one web-based system.",
    ]:
        add_bullet(document, item)

    add_callout(
        document,
        "Careful wording",
        "Do not say all features work fully offline. The interface and cached pages may load, but live data, login, registrations, and database actions still need internet.",
    )


def add_usage_section(document: Document) -> None:
    document.add_heading("2. How Users Can Use It", level=1)
    document.add_paragraph(
        "There are two valid ways to use the site. This is the easiest way to explain it without making installation sound required."
    )
    for item in [
        "Use in browser: Open the website link and use LYDO Connect normally.",
        "Install as PWA: Add it to the home screen or install it from the browser for faster repeat access.",
        "For one-time users: Browser use is enough.",
        "For frequent users: Installing is more convenient because the system is one tap away.",
    ]:
        add_bullet(document, item)

    document.add_heading("Typical installed PWA experience", level=2)
    for item in [
        "The user sees a LYDO Connect icon on the device.",
        "The user taps the icon instead of typing the URL.",
        "The site opens like an app.",
        "The user still receives the same web system and the same updated content when online.",
    ]:
        add_number(document, item)

    document.add_heading("Optional Demo Flow", level=2)
    for item in [
        "Open the deployed LYDO Connect website in a supported browser.",
        "Use the browser install option, usually named Install, Add to Home Screen, or Create shortcut.",
        "After installation, show that a LYDO Connect icon appears on the device.",
        "Open the icon and explain that it still uses the same web system, just in an app-like form.",
        "Clarify that this step is optional and mainly useful for repeat users.",
    ]:
        add_number(document, item)


def add_panel_scripts(document: Document) -> None:
    document.add_heading("3. Ready-to-Say Panel Scripts", level=1)
    add_script_block(
        document,
        "3.1 Main Explanation",
        "Our system uses PWA, or Progressive Web App. Ibig sabihin po, kahit website siya, pwede po siyang gamitin na parang mobile application. The user can open it in the browser, install it on their device, and access it directly from the home screen. This helps users access the system faster and more conveniently without downloading from an app store.",
    )
    add_script_block(
        document,
        "3.2 Shorter Version",
        "PWA po siya kasi binibigyan namin ang user ng choice: gamitin as website, or gamitin like an app. Hindi siya required, but it makes the system more practical for users who access it often.",
    )
    add_script_block(
        document,
        "3.3 Super Simple Version",
        "PWA po is like giving our website a costume of an app. Website pa rin siya, pero kapag ginamit ng user, feeling niya app ang gamit niya.",
    )
    add_script_block(
        document,
        "3.4 One-Line Defense",
        "We are not saying users must install it. We are saying the system is capable of being installed and used like an app.",
    )


def add_browser_vs_install(document: Document) -> None:
    document.add_heading("4. If Asked: Browser naman pwede, why install pa?", level=1)
    add_callout(
        document,
        "Direct answer",
        "Yes po, pwede naman talaga gamitin sa browser. Hindi po required i-install ang PWA. Optional lang po siya.",
    )
    document.add_paragraph(
        "The installation option is mainly for convenience. If a user often uses LYDO Connect, an installed PWA makes the system easier to access because the icon is already on the home screen."
    )
    for item in [
        "No need to type the website link again.",
        "No need to search browser history or bookmarks.",
        "One tap access, like a regular app.",
        "Useful for admins, staff, or youth users who repeatedly access the system.",
        "For one-time or occasional users, the browser is enough.",
    ]:
        add_bullet(document, item)

    add_callout(
        document,
        "Example",
        "Parang YouTube or Facebook. Pwede naman sa browser, pero mas convenient kapag may app or icon kasi one tap lang.",
        fill="F4F6F9",
    )
    add_script_block(
        document,
        "4.1 Panel Answer",
        "Hindi po mandatory ang installation. The system still works in the browser. PWA installation is just an added convenience for users who frequently use the system. It gives them quick access from their home screen and a more app-like experience without needing to download from Play Store or App Store.",
    )


def add_why_emphasize(document: Document) -> None:
    document.add_heading("5. If Asked: If optional, why emphasize PWA?", level=1)
    add_callout(
        document,
        "Direct answer",
        "We emphasize PWA not because installation is required, but because it shows the system is designed to be more accessible, flexible, and app-like when needed.",
    )
    document.add_paragraph("The important point is choice. PWA means the system supports both casual and frequent use.")
    for item in [
        "Casual users can open it as a normal website.",
        "Frequent users can install it for faster access.",
        "Admins and staff can treat it like a working tool on their device.",
        "Users do not need to download a separate Android or iOS app.",
        "The same system can serve phones, tablets, and computers.",
        "It shows the project is prepared for real-world use, not just classroom demonstration.",
    ]:
        add_bullet(document, item)

    add_script_block(
        document,
        "5.1 Panel Answer",
        "Although installation is optional, we emphasize PWA because it gives users flexibility. They can use it as a normal website, or install it if they want faster and easier access. So PWA is not a requirement, but an enhancement that improves accessibility, convenience, and usability.",
    )
    add_script_block(
        document,
        "5.2 Short Answer",
        "PWA po siya kasi binibigyan namin ang user ng choice: gamitin as website, or gamitin like an app. Hindi siya required, but it makes the system more practical for users who access it often.",
    )


def add_site_specific_notes(document: Document) -> None:
    document.add_heading("6. Site-Specific Notes for LYDO Connect", level=1)
    document.add_paragraph(
        "Based on the current project setup, LYDO Connect includes web app manifests and a production service worker."
    )
    for item in [
        "The public site manifest is named manifest.webmanifest and uses display: standalone.",
        "The admin portal manifest is named manifest-admin.webmanifest and starts at /admin.",
        "The service worker is registered only in production mode, so PWA behavior is expected after build/deployment.",
        "The service worker caches the app shell and same-origin GET requests, helping the interface and cached pages load faster.",
        "Because the system uses online data and authentication, database-dependent actions still need internet.",
    ]:
        add_bullet(document, item)

    add_callout(
        document,
        "Safe demo wording",
        "Our PWA support improves access and repeat use. It does not replace the need for internet when users log in, register, sync data, or load live records.",
    )


def add_qa_table(document: Document) -> None:
    document.add_heading("7. Quick Panelist Q&A", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3100, 6260])
    headers = table.rows[0].cells
    headers[0].text = "Possible Question"
    headers[1].text = "Recommended Answer"

    rows = [
        (
            "What is PWA?",
            "A website that can behave like an app. It can be opened in the browser or installed for app-like access.",
        ),
        (
            "Is installation required?",
            "No. Browser use is still valid. Installation is optional and mainly for convenience.",
        ),
        (
            "Why need install if browser works?",
            "For frequent users, installation gives one tap access from the home screen and a more app-like experience.",
        ),
        (
            "Why emphasize PWA if optional?",
            "Because it shows flexibility and accessibility. The system supports both website use and app-like use.",
        ),
        (
            "Does it work offline?",
            "Some cached parts or the app shell can load, but live data, login, registration, and database actions still need internet.",
        ),
        (
            "Why not make a separate mobile app?",
            "A PWA lets one system work across devices without separate Play Store or App Store deployment.",
        ),
    ]

    for question, answer in rows:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer
    format_table_text(table, header_rows=1)


def add_closing_summary(document: Document) -> None:
    document.add_heading("8. Final Summary to Memorize", level=1)
    add_callout(
        document,
        "Memorize this",
        "PWA is not about forcing users to install the system. It is about giving them an option to use the website like an app when that is more convenient.",
    )
    document.add_paragraph(
        "Best final line: LYDO Connect remains accessible through the browser, but its PWA support makes it more practical for repeat users because they can install it, open it from the home screen, and use a more app-like experience without downloading a separate mobile app."
    )

    closing = document.add_paragraph("End of Document")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(12)
    run = closing.runs[0]
    set_font(run, size=10, italic=True, color=MUTED)


def build_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_styles(document)
    set_footer(document)
    add_title_block(document)
    add_simple_pwa_explanation(document)
    add_usage_section(document)
    add_panel_scripts(document)
    add_browser_vs_install(document)
    add_why_emphasize(document)
    add_site_specific_notes(document)
    add_qa_table(document)
    add_closing_summary(document)
    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    saved = Document(OUTPUT_PATH)
    heading_count = sum(1 for paragraph in saved.paragraphs if paragraph.style.name.startswith("Heading"))
    table_count = len(saved.tables)
    print(OUTPUT_PATH.resolve())
    print(f"Verified readable DOCX: {heading_count} headings, {table_count} tables.")


if __name__ == "__main__":
    main()
