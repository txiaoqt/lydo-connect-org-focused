from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(46, 116, 181)
    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(7)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(46, 116, 181)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(31, 77, 120)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("LYDO CONNECT\nWHOLE-SITE USERFLOW GUIDE")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(26, 63, 122)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run("Separated Userflows for Youth and Admin")
    sub_run.font.name = "Calibri"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    sub_run.font.size = Pt(12)
    sub_run.italic = True

    document.add_paragraph()

    meta_table = document.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    meta_table.columns[0].width = Inches(2.0)
    meta_table.columns[1].width = Inches(4.5)
    metadata = [
        ("Document Purpose", "I-define ang complete userflow ng buong site, hiwalay para sa Youth at Admin."),
        ("Actors", "1) Youth Users  2) Admin Users"),
        ("Coverage", "Public pages, authentication, core modules, transparency, administration, and analytics."),
        ("Design Basis", "Current route map and module behavior in the existing LYDO Connect codebase."),
        ("Primary Output", "Operational flow guide for implementation, QA, demos, and onboarding."),
    ]
    for row, (label, value) in zip(meta_table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_number(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_step_block(document: Document, title: str, actions: list[str], outcomes: list[str]) -> None:
    document.add_heading(title, level=3)
    document.add_paragraph("User Actions:")
    for item in actions:
        add_bullet(document, item)
    document.add_paragraph("System Outcomes:")
    for item in outcomes:
        add_bullet(document, item)


def add_matrix_table(document: Document) -> None:
    document.add_heading("5. Page-to-User Responsibility Matrix", level=1)
    document.add_paragraph(
        "Quick reference kung sino ang primary actor sa bawat module at ano ang typical actions."
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Module / Page"
    headers[1].text = "Youth Role"
    headers[2].text = "Admin Role"
    headers[3].text = "Primary Purpose"

    rows = [
        ("Home, About, FAQs, Contacts, Site Map", "View and navigate information", "Indirect (content direction only)", "Public orientation and discovery"),
        ("Programs and Events (Public)", "Search, open details, register, unregister", "Publishes records through admin modules", "Participation funnel"),
        ("Event/Program Record Detail", "Complete registration form", "Controls availability/status/links", "Transaction point for participation"),
        ("Profile", "Update profile, review joined records", "Indirect via user governance tools", "Self-service account management"),
        ("Organizations (Public)", "Browse organizations and view details", "Maintains organization records", "Community partner visibility"),
        ("Transparency Reports", "Search and export public disclosures", "Uploads and maintains documents", "Transparency and compliance access"),
        ("Transparency Board and Financial Disclosure", "Monitor compliance and budget utilization", "Encodes board and financial data", "Public accountability"),
        ("Youth Desk (Public)", "Submit and track tickets", "Updates ticket status and handling notes", "Youth request loop"),
        ("Admin Dashboard", "Not accessible", "Monitor KPIs, filters, quick actions", "Control center"),
        ("Programs / Events Admin", "Not accessible", "Create, edit, filter, archive records", "Program lifecycle control"),
        ("Registrations Admin", "Not accessible", "Review participants, sync, export CSV", "Attendance and data ops"),
        ("Organizations Admin", "Not accessible", "Create/edit/delete organization profiles", "Stakeholder management"),
        ("Barangay Map Data Admin", "Not accessible", "Manage barangays and resident profile details", "Location and demographics data quality"),
        ("Transparency Docs Admin", "Not accessible", "Upload files and metadata", "Disclosure registry maintenance"),
        ("Transparency Board Config Admin", "Not accessible", "Maintain quarterly/monthly compliance rows", "Compliance publishing inputs"),
        ("Financial DSS Admin", "Not accessible", "Maintain rows, budgets, exports", "Financial analytics and reporting"),
        ("Youth Desk Admin", "Not accessible", "Filter tickets and update statuses", "Service response management"),
        ("Users / Roles / Audit Logs", "Not accessible", "Govern access and trace changes", "Security and accountability"),
        ("Trends and Analytics", "Not accessible", "Analyze outcomes, encode result records, export reports", "Decision intelligence"),
    ]

    for module, youth_role, admin_role, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = module
        cells[1].text = youth_role
        cells[2].text = admin_role
        cells[3].text = purpose


def add_tagged_non_negotiables_section(document: Document) -> None:
    document.add_heading("7. Tagged Non-Negotiables in Userflow", level=1)
    document.add_paragraph(
        "This section classifies core flow checkpoints using required tags so teams can quickly distinguish strict gates versus optional paths."
    )

    document.add_heading("7.1 Tag Legend", level=2)
    add_bullet(
        document,
        "[NON-NEGOTIABLE] Required gate. If not satisfied, the user cannot proceed in that target flow.",
    )
    add_bullet(
        document,
        "[CONDITIONAL] Required only when a user enters a specific protected or transactional path.",
    )
    add_bullet(
        document,
        "[NEGOTIABLE] Optional path. User may skip this and still continue through other valid routes.",
    )

    document.add_heading("7.2 Youth Flow Tagged Rules", level=2)
    youth_table = document.add_table(rows=1, cols=3)
    youth_table.style = "Table Grid"
    youth_headers = youth_table.rows[0].cells
    youth_headers[0].text = "Userflow Item"
    youth_headers[1].text = "Tag"
    youth_headers[2].text = "Reason / System Behavior"

    youth_rows = [
        (
            "Browse public pages (Home, Programs, Events, Orgs, Transparency)",
            "[NEGOTIABLE]",
            "Sign-in is not required for viewing public content.",
        ),
        (
            "Create an account (Sign Up)",
            "[NEGOTIABLE]",
            "Account creation is optional for browsing-only users.",
        ),
        (
            "Sign in to register for a program/event",
            "[NON-NEGOTIABLE]",
            "Registration flow blocks unauthenticated users and shows sign-in requirement.",
        ),
        (
            "Complete registration form validation before submit",
            "[NON-NEGOTIABLE]",
            "Form errors stop registration submission until corrected.",
        ),
        (
            "Sign in to submit Youth Desk ticket",
            "[NON-NEGOTIABLE]",
            "Ticket submission button and logic require authenticated user session.",
        ),
        (
            "Sign in to open Profile page and manage joined records",
            "[NON-NEGOTIABLE]",
            "Profile route redirects unauthenticated users to sign-in.",
        ),
        (
            "Accept active Terms and Privacy policy after login",
            "[CONDITIONAL]",
            "Applies to authenticated non-admin users when policy acceptance is pending.",
        ),
        (
            "Use external registration/source links",
            "[NEGOTIABLE]",
            "External links are supplementary; local page details remain accessible.",
        ),
    ]

    for item, tag, reason in youth_rows:
        cells = youth_table.add_row().cells
        cells[0].text = item
        cells[1].text = tag
        cells[2].text = reason

    document.add_heading("7.3 Admin Flow Tagged Rules", level=2)
    admin_table = document.add_table(rows=1, cols=3)
    admin_table.style = "Table Grid"
    admin_headers = admin_table.rows[0].cells
    admin_headers[0].text = "Userflow Item"
    admin_headers[1].text = "Tag"
    admin_headers[2].text = "Reason / System Behavior"

    admin_rows = [
        (
            "Sign in as admin to access /admin portal",
            "[NON-NEGOTIABLE]",
            "Admin route is guard-protected; non-admin sessions are redirected.",
        ),
        (
            "Use valid admin role/session for protected modules",
            "[NON-NEGOTIABLE]",
            "RequireAdmin guard enforces admin-only access for portal workflows.",
        ),
        (
            "Maintain core records (Programs, Events, Organizations, Barangays)",
            "[NON-NEGOTIABLE]",
            "These records are foundational inputs for youth-facing pages and reports.",
        ),
        (
            "Run registration sync and exports",
            "[CONDITIONAL]",
            "Required for operations that rely on external attendance ingestion/reporting windows.",
        ),
        (
            "Encode outcome records in Trends and Analytics",
            "[CONDITIONAL]",
            "Required when formal outcome documentation/reporting is expected.",
        ),
        (
            "Use Users, Roles, and Audit Logs governance modules",
            "[CONDITIONAL]",
            "Becomes required during access changes, incident review, or compliance checks.",
        ),
        (
            "Open dashboard quick actions/shortcuts",
            "[NEGOTIABLE]",
            "Shortcuts speed up work but are not mandatory for completing admin tasks.",
        ),
    ]

    for item, tag, reason in admin_rows:
        cells = admin_table.add_row().cells
        cells[0].text = item
        cells[1].text = tag
        cells[2].text = reason

    document.add_heading("7.4 Direct Clarification on Login Example", level=2)
    document.add_paragraph(
        "Your interpretation is correct: for Youth users, login is [NEGOTIABLE] for browsing public content, but it becomes [NON-NEGOTIABLE] for protected actions like registration, profile access, and Youth Desk submission."
    )


def build_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_styles(document)
    add_title_block(document)

    document.add_heading("1. System Context and Role Separation", level=1)
    document.add_paragraph(
        "Ang LYDO Connect ay may dalawang pangunahing user groups na may malinaw na separation ng access at responsibilities: Youth users at Admin users."
    )
    add_bullet(document, "Youth users operate in public/youth-facing routes (home, programs, events, organizations, transparency, profile, Youth Desk).")
    add_bullet(document, "Admin users operate inside the Admin Portal (/admin) with role-protected access.")
    add_bullet(document, "Non-admin users cannot open protected admin pages; admins are redirected away from public-only workflows when needed.")
    add_bullet(document, "Authenticated youth users are required to accept active Terms and Privacy policy before continuing.")

    document.add_heading("1.1 Tag Legend Used in This Document", level=2)
    add_bullet(document, "[NON-NEGOTIABLE] Mandatory gate for the target flow.")
    add_bullet(document, "[CONDITIONAL] Mandatory only for specific scenarios or protected paths.")
    add_bullet(document, "[NEGOTIABLE] Optional path; users can continue through alternate valid routes.")

    document.add_heading("2. Youth Userflow (End-to-End)", level=1)
    document.add_paragraph(
        "Ito ang primary journey ng youth participant mula discovery hanggang participation at feedback."
    )

    add_step_block(
        document,
        "2.1 Entry and Public Discovery",
        [
            "Open home page and use navbar to explore Programs, Events, Organizations, and Transparency pages.",
            "Review About, FAQs, Contacts, and legal pages for context and trust.",
            "Use search/filter controls on listing pages before choosing a specific record.",
        ],
        [
            "Site remains fully browsable kahit guest pa lang ang user.",
            "Each listing page narrows records based on user-selected filters and search terms.",
            "User can proceed to account creation or direct sign in once ready to transact.",
        ],
    )

    add_step_block(
        document,
        "2.2 Account Creation (Youth Sign Up)",
        [
            "Go to Sign Up and fill full name, email, contact number, barangay, and password.",
            "Confirm form details and explicitly agree to Terms and Privacy policy in the confirmation step.",
            "Submit account creation request.",
        ],
        [
            "Validation checks are applied (required fields, format checks, and policy agreement).",
            "Account is created, then user is signed out intentionally to enforce manual sign-in flow.",
            "User is redirected to Sign In page for authenticated access.",
        ],
    )

    add_step_block(
        document,
        "2.3 Sign In and Policy Gate",
        [
            "Open Sign In and authenticate using youth account credentials.",
            "If active policy version requires acceptance, review modal and accept before proceeding.",
            "Continue to browsing or profile after successful authentication and policy agreement.",
        ],
        [
            "Session becomes active with youth role context.",
            "Policy agreement state is checked automatically for authenticated non-admin users.",
            "Declining policy signs user out; acceptance allows normal use of protected youth actions.",
        ],
    )

    add_step_block(
        document,
        "2.4 Program/Event Participation Flow",
        [
            "Open Programs or Events list, then choose a specific record.",
            "On record page, review details tab then switch to registration tab.",
            "Submit registration form (or use external form path where enabled).",
            "Optionally cancel/unregister later via the same record or profile page.",
        ],
        [
            "Registration requires authenticated user identity.",
            "Submitted registration is stored and reflected in joined status indicators.",
            "Joined entries appear in user profile and become visible in admin registrations/analytics views.",
        ],
    )

    add_step_block(
        document,
        "2.5 Personal Profile and Participation History",
        [
            "Open Profile to review joined events/programs and account details.",
            "Update editable profile settings and save.",
            "Leave joined program/event when user opts out.",
        ],
        [
            "Profile settings persist to backend profile records.",
            "Joined lists provide personal activity history context.",
            "Leave action updates registration state and downstream analytics totals.",
        ],
    )

    add_step_block(
        document,
        "2.6 Transparency and Youth Service Flow",
        [
            "Open transparency pages to check disclosures, board status, financial views, and barangay map data.",
            "Use report filters and export options in disclosure registry.",
            "Open Youth Desk to submit concerns/requests and monitor personal ticket statuses.",
        ],
        [
            "Transparency pages expose public accountability information for youth/public viewers.",
            "Youth Desk creates trackable tickets with status lifecycle (received to closed).",
            "Youth can monitor ticket updates in their own My Tickets area.",
        ],
    )

    document.add_heading("2.7 Youth Flow - Common Decision Branches", level=2)
    add_bullet(document, "If guest attempts a transaction (register/submit ticket), system prompts sign-in first.")
    add_bullet(document, "If required policy data cannot be loaded, user gets retry/sign-out fallback.")
    add_bullet(document, "If no matching records are found after filters, user receives empty-state guidance and can clear filters.")

    document.add_heading("3. Admin Userflow (End-to-End)", level=1)
    document.add_paragraph(
        "Ito ang operational journey ng admin team mula data setup, monitoring, intervention, hanggang reporting."
    )

    add_step_block(
        document,
        "3.1 Admin Access and Authentication",
        [
            "Open admin sign-in route and authenticate as admin.",
            "Enter Admin Portal and land on Dashboard.",
            "Use sidebar sections to move across operational modules.",
        ],
        [
            "Admin-only guard protects /admin routes.",
            "Role-based redirection prevents non-admin access.",
            "Sidebar state (including collapse preference) is retained for workflow continuity.",
        ],
    )

    add_step_block(
        document,
        "3.2 Dashboard Control Center",
        [
            "Review KPI cards, quick actions, and search results.",
            "Use dashboard filters (date range, role, barangay, municipality, status filters).",
            "Jump directly to sections (Programs, Events, Registrations, Trends and Analytics, etc.).",
        ],
        [
            "Dashboard aggregates multi-module snapshots in one view.",
            "Filter state narrows visible insights without leaving dashboard.",
            "Quick actions reduce clicks for high-frequency admin tasks.",
        ],
    )

    add_step_block(
        document,
        "3.3 Youth Records Operations",
        [
            "Programs Management: create/edit/delete program entries, manage status, links, locations, and sync settings.",
            "Events Management: create/edit/delete event entries with similar controls.",
            "Registrations: monitor participant rows, run sync now, and export filtered CSV.",
            "Organizations: maintain stakeholder records with profile details and status.",
            "Barangay Map Data: manage barangay records and update resident profile information.",
        ],
        [
            "Youth-facing data quality is controlled centrally by admin.",
            "Registration and sync tools keep attendance datasets operational.",
            "Updates in these modules directly affect what youth users can view and join.",
        ],
    )

    add_step_block(
        document,
        "3.4 Transparency Operations",
        [
            "Transparency Docs: upload/manage disclosure documents and metadata.",
            "Transparency Board Config: maintain board rows and monthly compliance rows.",
            "Financial DSS: maintain financial rows/budgets, review charts, export Excel/PDF.",
            "Youth Desk Admin: filter tickets, open details, update status and handling notes.",
        ],
        [
            "Public transparency pages are fed by these admin-maintained records.",
            "Financial and compliance data become traceable and report-ready.",
            "Youth requests receive lifecycle handling from received to closed.",
        ],
    )

    add_step_block(
        document,
        "3.5 Governance and Security Controls",
        [
            "Users: search/filter users, edit profiles, and delete user records when required.",
            "Roles and Permissions: create/edit/delete role entries and update access definitions.",
            "Audit Logs: trace who changed what, when, and in which data section.",
        ],
        [
            "Access governance remains centralized and reviewable.",
            "Audit history supports accountability, incident investigation, and compliance checks.",
            "User/role updates affect effective capability boundaries across the portal.",
        ],
    )

    add_step_block(
        document,
        "3.6 Trends and Analytics Operational Loop",
        [
            "Open Trends and Analytics section under dedicated sidebar group.",
            "Apply filters (period, year, quarter, month, sector, barangay, record type, attendance, status).",
            "Review KPI cards, trend charts, sector/barangay views, alerts, and performance lists.",
            "Create/edit/delete program outcome records and export summary/participation/outcome datasets.",
            "Use print/export outputs for meetings, planning sessions, and compliance reporting.",
        ],
        [
            "Admin gets consolidated evidence for performance and reach decisions.",
            "Outcome records bridge operational activity with transparency documentation.",
            "Filtered exports make reporting context-specific and reproducible.",
        ],
    )

    document.add_heading("3.7 Admin Flow - Common Decision Branches", level=2)
    add_bullet(document, "If required schema/table is missing (ex: outcomes table), page warns and keeps safe fallback behavior where possible.")
    add_bullet(document, "If sync or load actions fail, admin can retry and review diagnostics/alerts.")
    add_bullet(document, "If no filtered rows match, module shows empty state and encourages filter reset or broadened criteria.")

    document.add_heading("4. Cross-Role Lifecycle (Whole Site)", level=1)
    document.add_paragraph(
        "Pinagdurugtong ng flow na ito ang youth-side participation at admin-side operations."
    )
    lifecycle_steps = [
        "Admin creates or updates Programs/Events with status, schedule, and location.",
        "Youth discovers records from public pages and submits registrations.",
        "Registrations feed admin Registrations module and analytics datasets.",
        "Admin monitors participation quality, sync health, and low-performing records.",
        "Admin encodes outcome records and uploads supporting transparency documents.",
        "Public users view updated transparency outputs and service status via Youth Desk.",
        "Governance modules (Users, Roles, Audit Logs) ensure continuity and accountability.",
    ]
    for step in lifecycle_steps:
        add_number(document, step)

    add_matrix_table(document)

    document.add_heading("6. Condensed Role-Based Flow Map", level=1)
    document.add_paragraph("Youth flow (simplified):")
    youth_flow = "Discover -> Sign Up/Sign In -> Accept Policy -> Browse/Join Programs and Events -> Manage Profile -> Submit/Track Youth Desk Ticket -> Repeat Participation"
    add_bullet(document, youth_flow)
    document.add_paragraph("Admin flow (simplified):")
    admin_flow = "Admin Sign In -> Dashboard -> Maintain Core Records -> Monitor Registrations and Sync -> Manage Transparency and Financial Data -> Run Trends and Analytics -> Govern Users/Roles/Audit -> Report and Iterate"
    add_bullet(document, admin_flow)

    add_tagged_non_negotiables_section(document)

    document.add_heading("8. Purpose and Usage", level=1)
    document.add_paragraph(
        "This userflow guide is intended for implementation planning, QA test case design, onboarding, stakeholder demos, and process alignment. "
        "By separating Youth and Admin flows, mas malinaw ang role boundaries, expected actions, and system responsibilities for each side of the LYDO Connect platform."
    )

    document.add_paragraph()
    closing = document.add_paragraph("End of Document")
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    closing.runs[0].italic = True

    return document


def main() -> None:
    output_path = Path("LYDO_Connect_Whole_Site_Userflows_Youth_and_Admin.docx")
    document = build_document()
    document.save(output_path)
    print(str(output_path.resolve()))


if __name__ == "__main__":
    main()


