from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(46, 116, 181)
    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(7)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(46, 116, 181)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(31, 77, 120)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("LYDO CONNECT ADMIN PORTAL\nTRENDS AND ANALYTICS")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(26, 63, 122)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run("Feature Explanation and Functional Guide for Programs and Events")
    sub_run.font.name = "Calibri"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    sub_run.font.size = Pt(12)
    sub_run.italic = True

    document.add_paragraph()

    meta_table = document.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_table.columns[0].width = Inches(1.9)
    meta_table.columns[1].width = Inches(4.6)
    metadata = [
        ("Document Purpose", "Ipaliwanag ang lahat ng sections ng Trends and Analytics page."),
        ("Coverage", "Programs and Events trends, KPI logic, charts, alerts, exports, and outcome records."),
        ("Target Audience", "LYDO Staff, SK, and Admin users."),
        ("System Context", "LYDO Connect Admin Portal."),
    ]
    for row, (label, value) in zip(meta_table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_number(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_section_detail(
    document: Document,
    title: str,
    purpose: str,
    how_it_works: list[str],
    business_value: list[str],
) -> None:
    document.add_heading(title, level=2)
    document.add_paragraph(f"Purpose: {purpose}")

    document.add_heading("How it Functions", level=3)
    for item in how_it_works:
        add_bullet(document, item)

    document.add_heading("Why it Matters", level=3)
    for item in business_value:
        add_bullet(document, item)


def build_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_styles(document)
    add_title_block(document)

    document.add_heading("1. Overview", level=1)
    document.add_paragraph(
        "Ang Trends and Analytics page ay dedicated section para sa monitoring ng program at event outcomes. "
        "Pinagsasama nito ang participation records, attendance behavior, barangay reach, integration health, at outcome documentation."
    )
    document.add_paragraph(
        "Main goal: tulungan ang admin team gumawa ng data-driven decisions kung ano ang epektibong activities, alin ang kailangan i-improve, "
        "at alin ang nangangailangan ng immediate action."
    )

    document.add_heading("2. Data Sources Used by the Page", level=1)
    source_items = [
        "Programs table: title, sector, status, barangay, dates.",
        "Events table: title, sector, status, barangay, event dates.",
        "Program registrations and Event registrations: participant and attendance lifecycle.",
        "Barangays table: barangay names and youth population baseline for reach percentage.",
        "Disclosure documents table (program_outcome type): linked outcome reports for transparency.",
        "Program outcomes table: target vs actual, completion percentage, summary, challenges, recommendations.",
    ]
    for source in source_items:
        add_bullet(document, source)
    document.add_paragraph(
        "If some columns are missing (example: sync columns), the page uses safe fallback defaults. "
        "If the program_outcomes table is not yet migrated, the page shows an explicit warning and keeps read-only analytics working."
    )

    document.add_heading("3. Global Controls and Filters", level=1)
    document.add_paragraph(
        "Ang filter panel ang core control layer ng page. Lahat ng cards, charts, tables, at alerts ay naka-base sa active filter state."
    )
    filters = [
        "Date Range: All time, last 30/90/180/365 days.",
        "Year, Quarter, and Month: Period slicing para sa reporting windows.",
        "Sector: Focus by sector (Education, Health, etc.).",
        "Barangay: Focus by local area coverage.",
        "Record Type: Programs only, Events only, or both.",
        "Attendance Status: registered, waitlisted, attended, no_show, cancelled.",
        "Program Status: upcoming, ongoing, past, archived, etc.",
        "Trend Granularity: monthly, quarterly, yearly aggregation sa trend chart.",
    ]
    for item in filters:
        add_bullet(document, item)
    document.add_paragraph(
        "Buttons available: Refresh Data, Filter toggle, Clear Filters, Export Summary, Export Participation CSV, and Print/PDF."
    )

    document.add_heading("4. KPI Cards and Exact Logic", level=1)
    add_bullet(document, "Programs Conducted: count ng filtered programs.")
    add_bullet(document, "Registered Participants: registrations excluding cancelled.")
    add_bullet(document, "Attended Participants: rows with status attended.")
    add_bullet(document, "Attendance Rate: attended / (attended + registered + waitlisted + no_show).")
    add_bullet(document, "No-show / Cancelled: count of no_show plus cancelled rows.")
    add_bullet(document, "Unique Youth Reached: unique email set from active registrations.")
    add_bullet(document, "Active Barangays Reached: unique barangay IDs from active registrations.")
    add_bullet(document, "Outcome Reports Uploaded: count of disclosure docs tagged as program_outcome.")

    document.add_heading("5. Detailed Feature-by-Feature Explanation", level=1)

    add_section_detail(
        document,
        "5.1 Participation Trends Chart",
        "Makita ang growth or decline ng registrations at attendance over time.",
        [
            "Nag-aaggregate ang system by selected granularity (month/quarter/year).",
            "Series included: Registered, Attended, and No Show.",
            "May auto-computed indicator na percentage growth vs previous period.",
        ],
        [
            "Quickly identifies seasonality at campaign impact.",
            "Supports planning for schedule, manpower, and promotion timing.",
        ],
    )

    add_section_detail(
        document,
        "5.2 Attendance Status Breakdown",
        "Makita ang distribution ng participant states.",
        [
            "Pie chart slices are grouped by registration_status values.",
            "Statuses include attended, registered, waitlisted, no_show, cancelled based on filtered rows.",
        ],
        [
            "Helps spot behavior pattern, especially high no_show rates.",
            "Supports policy updates for reminders and attendance enforcement.",
        ],
    )

    add_section_detail(
        document,
        "5.3 Sector Performance",
        "I-compare ang program outcomes by sector.",
        [
            "For each sector: total registrations, total attended, computed attendance rate.",
            "Program count per sector is included para makita participation volume vs supply of activities.",
        ],
        [
            "Shows strongest and weakest sectors.",
            "Guides budget and program design per sector.",
        ],
    )

    add_section_detail(
        document,
        "5.4 Barangay Reach (Chart + Table)",
        "I-track kung gaano kalawak ang local youth engagement per barangay.",
        [
            "Computes unique participants per barangay from active registration data.",
            "Computes reach percent = unique participants / youth population.",
            "Tracks programs conducted and attendance rate per barangay.",
            "Table includes visual progress bar for reach percentage.",
        ],
        [
            "Identifies underserved barangays.",
            "Supports location targeting and outreach planning.",
        ],
    )

    add_section_detail(
        document,
        "5.5 Top Performing Programs",
        "I-rank ang programs with best attendance performance.",
        [
            "Filtered to programs with actual participation records.",
            "Sorted by attendance rate, then attended volume.",
            "Shows if an outcome record already exists or still missing.",
        ],
        [
            "Best-practice identification.",
            "Reusable template for future program design.",
        ],
    )

    add_section_detail(
        document,
        "5.6 Low Performing Programs",
        "I-highlight ang programs with weak attendance behavior.",
        [
            "Focuses on programs with meaningful sample size (minimum registrations threshold).",
            "Sorted by low attendance rate and high no_show count.",
        ],
        [
            "Flags programs needing corrective intervention.",
            "Prevents recurring low-impact activities.",
        ],
    )

    add_section_detail(
        document,
        "5.7 Operational Alerts",
        "Automated early-warning panel for actionable risks.",
        [
            "Programs Missing Outcome Records: completed programs without encoded outcomes.",
            "Low Attendance Programs: attendance rate below threshold.",
            "Failed Attendance Sync: records that failed integration sync.",
            "Programs with No Attendance Data: ended programs with zero registrations.",
            "Low Barangay Reach: barangays below baseline reach percentage.",
        ],
        [
            "Speeds up admin response time.",
            "Makes issues visible even before manual report generation.",
        ],
    )

    add_section_detail(
        document,
        "5.8 Attendance Sync Diagnostics",
        "I-monitor integration health ng attendance pipeline.",
        [
            "Bar chart for sync states: pending, failed, synced, skipped.",
            "Shows exact count of rows per state under current filters.",
            "Can be exported as CSV for technical follow-up.",
        ],
        [
            "Improves reliability of external form integrations.",
            "Provides traceable evidence for sync troubleshooting.",
        ],
    )

    add_section_detail(
        document,
        "5.9 Registration Source Mix",
        "Makita ang origin ng registration data.",
        [
            "Pie chart groups source values such as portal_direct, admin_csv_sync, imported.",
            "Distribution updates live with active filters.",
        ],
        [
            "Helps validate if adoption is coming from portal or external channels.",
            "Supports governance and data quality checks.",
        ],
    )

    add_section_detail(
        document,
        "5.10 Program Outcome Records (CRUD)",
        "Structured repository ng measurable outcomes per program.",
        [
            "Create/Edit/Delete outcome rows directly from the page.",
            "Fields include: target participants, actual participants, completion percent, objectives achieved, outcome summary, challenges, recommendations.",
            "Supports optional link to transparency document (program outcome report).",
            "Completion percent can be encoded or derived from target vs actual.",
            "When table is not yet migrated, page displays warning and keeps analytics available.",
        ],
        [
            "Standardizes post-activity reporting.",
            "Connects operational outcomes with transparency and compliance artifacts.",
        ],
    )

    document.add_heading("6. Export and Reporting Functions", level=1)
    export_items = [
        "Export Summary CSV: high-level KPI snapshot.",
        "Export Participation CSV: row-level participation and attendance diagnostics.",
        "Export Outcomes CSV: outcome records for office reporting.",
        "Print/PDF: fast documentation for meetings and reviews.",
    ]
    for item in export_items:
        add_bullet(document, item)
    document.add_paragraph(
        "All exports follow the currently applied filters so generated files are context-specific."
    )

    document.add_heading("7. Responsive Behavior and Usability", level=1)
    add_bullet(document, "Grid-based layout adapts from desktop multi-column view to mobile single-column stack.")
    add_bullet(document, "Filter controls wrap and remain usable on smaller screens.")
    add_bullet(document, "Tables are horizontally scrollable on narrow viewports.")
    add_bullet(document, "Cards and chart blocks retain readable spacing and hierarchy.")

    document.add_heading("8. Intended Operational Workflow", level=1)
    workflow_steps = [
        "Open Trends and Analytics section.",
        "Apply period, sector, and barangay filters.",
        "Review KPI cards and participation trend chart.",
        "Check alerts and low-performing programs for intervention list.",
        "Encode or update Program Outcome Records.",
        "Export summary and details for formal reporting.",
    ]
    for step in workflow_steps:
        add_number(document, step)

    document.add_heading("9. Key Purpose Summary", level=1)
    document.add_paragraph(
        "Ang pangunahing purpose ng Trends and Analytics ay gawing measurable, traceable, at actionable ang outcomes ng LYDO programs at events. "
        "Ito ang central monitoring page kung saan nakikita ang participation trend, attendance quality, barangay penetration, integration health, at documented outcomes, "
        "para mas mabilis at mas tama ang planning at decision-making ng admin team."
    )

    document.add_paragraph()
    closing = document.add_paragraph("End of Document")
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    closing.runs[0].italic = True

    return document


def main() -> None:
    output_path = Path("Trends_and_Analytics_Feature_Explanation.docx")
    doc = build_document()
    doc.save(output_path)
    print(str(output_path.resolve()))


if __name__ == "__main__":
    main()

