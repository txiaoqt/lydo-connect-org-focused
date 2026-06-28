from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(22, 39, 61)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "EAF1FB"
ALT_FILL = "F7F9FC"
BORDER = "C9D4E5"


@dataclass
class DemoSection:
    number: str
    title: str
    goal: str
    actions: list[str]
    script_lines: list[str]
    key_points: list[str]


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def style_run(run, *, name: str = "Calibri", size: int = 11, bold: bool = False, italic: bool = False, color: RGBColor = INK) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_block(document: Document, title: str, subtitle: str, audience: str, outcome: str) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_paragraph, after=4)
    title_run = title_paragraph.add_run(title)
    style_run(title_run, size=22, bold=True, color=BLUE)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle_paragraph, after=12)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    style_run(subtitle_run, size=11, italic=True, color=GRAY)

    intro_table = document.add_table(rows=3, cols=2)
    intro_table.style = "Table Grid"
    intro_table.autofit = False
    labels = ["Audience", "Output Goal", "Tone"]
    values = [
        audience,
        outcome,
        "Taglish, presenter-friendly, and detailed enough for demo, defense, or onboarding use.",
    ]
    for row_index, row in enumerate(intro_table.rows):
        for column_index, width in enumerate((1.7, 4.8)):
            set_cell_width(row.cells[column_index], width)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        label_paragraph = row.cells[0].paragraphs[0]
        set_paragraph_spacing(label_paragraph, after=0)
        label_run = label_paragraph.add_run(labels[row_index])
        style_run(label_run, size=10, bold=True, color=DARK_BLUE)

        value_paragraph = row.cells[1].paragraphs[0]
        set_paragraph_spacing(value_paragraph, after=0)
        value_run = value_paragraph.add_run(values[row_index])
        style_run(value_run)

    document.add_paragraph()


def add_note_box(document: Document, title: str, lines: list[str]) -> None:
    table = document.add_table(rows=1 + len(lines), cols=1)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        set_cell_width(row.cells[0], 6.5)
    set_cell_shading(table.rows[0].cells[0], LIGHT_FILL)
    heading_paragraph = table.rows[0].cells[0].paragraphs[0]
    set_paragraph_spacing(heading_paragraph, after=0)
    heading_run = heading_paragraph.add_run(title)
    style_run(heading_run, size=11, bold=True, color=DARK_BLUE)

    for index, line in enumerate(lines, start=1):
        paragraph = table.rows[index].cells[0].paragraphs[0]
        set_paragraph_spacing(paragraph, after=0)
        run = paragraph.add_run(line)
        style_run(run)

    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=2)
        run = paragraph.add_run(item)
        style_run(run)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=2)
        run = paragraph.add_run(item)
        style_run(run)


def add_coverage_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    document.add_heading("Coverage Summary", level=1)
    document.add_paragraph(
        "Ito ang mabilis na reference ng mga sections na sakop ng script at kung ano ang primary focus ng bawat bahagi."
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Section", "Module / Screen", "Primary Focus"]
    widths = (1.0, 2.3, 3.2)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, after=0)
        run = paragraph.add_run(header)
        style_run(run, size=10, bold=True, color=DARK_BLUE)
    for section, module, focus in rows:
        row = table.add_row()
        values = (section, module, focus)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0)
            run = paragraph.add_run(value)
            style_run(run)
    document.add_paragraph()


def add_demo_section(document: Document, section: DemoSection) -> None:
    document.add_heading(f"{section.number}. {section.title}", level=1)
    goal_paragraph = document.add_paragraph()
    set_paragraph_spacing(goal_paragraph, after=8)
    goal_label = goal_paragraph.add_run("Goal: ")
    style_run(goal_label, bold=True, color=DARK_BLUE)
    goal_text = goal_paragraph.add_run(section.goal)
    style_run(goal_text)

    document.add_heading("Suggested On-Screen Actions", level=2)
    add_numbered(document, section.actions)

    document.add_heading("Sample Presenter Script", level=2)
    for line in section.script_lines:
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, after=4)
        run = paragraph.add_run(line)
        style_run(run)

    document.add_heading("Key Points to Emphasize", level=2)
    add_bullets(document, section.key_points)

    add_note_box(
        document,
        "Presenter Reminder",
        [
            "Panatilihing malinaw ang transition line bago lumipat sa susunod na screen.",
            "Kapag may loading state, revision state, o approval state, i-explain kung ano ang ibig sabihin nito sa actual workflow.",
        ],
    )


def add_closing_section(document: Document, lines: list[str]) -> None:
    document.add_heading("Closing Script", level=1)
    for line in lines:
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, after=4)
        run = paragraph.add_run(line)
        style_run(run)


def build_user_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_styles(document)

    add_title_block(
        document,
        "LYDO Connect User Side Demo Script",
        "Comprehensive presenter script for the organization and youth-facing user portal",
        "Student presenters, panel defense speakers, demo facilitators, and onboarding leads",
        "To explain the end-to-end user workflow from dashboard setup up to compliance, funding, liquidation, updates, and incentives.",
    )

    add_note_box(
        document,
        "How to Use This Script",
        [
            "Pwede itong basahin almost word-for-word habang nagde-demo, o gamitin bilang speaking guide kung gusto mo ng mas natural na delivery.",
            "Ang flow na ito ay naka-base sa current user portal modules: Dashboard, Organization Profile, Document Submissions, Budget Requests, Liquidation Reports, YPOP Incentive, News Releases, at Notifications.",
            "Kung may kulang na sample data sa demo environment, gamitin ang transition lines at expected outcome statements para maipaliwanag pa rin ang workflow.",
        ],
    )

    add_coverage_table(
        document,
        [
            ("1", "Dashboard", "Overall portal orientation and workflow status"),
            ("2", "Organization Profile", "Initial compliance setup and profile verification"),
            ("3", "Document Submissions", "Uploading, OCR review, and admin approval flow"),
            ("4", "Budget Requests", "Funding request creation, attachment, and review states"),
            ("5", "Liquidation Reports", "Post-activity reporting and liquidation follow-through"),
            ("6", "YPOP Incentive", "Semester-based incentive submission and qualification logic"),
            ("7", "News Releases", "User-facing updates and announcements"),
            ("8", "Notifications", "Action reminders and status awareness"),
        ],
    )

    sections = [
        DemoSection(
            number="1",
            title="Dashboard Orientation",
            goal="Ipakita agad kung paano naiintindihan ng user ang kabuuang status ng organization sa pagpasok pa lang sa portal.",
            actions=[
                "Mag-sign in bilang user at ipakita ang dashboard landing page.",
                "I-tour ang left navigation para makita ang Home, Compliance, Grants & Incentives, at News Releases.",
                "I-highlight ang progress or readiness cards para makita kung anong workflow ang tapos, kulang, o naka-pending pa.",
            ],
            script_lines=[
                "Sa user side, ito ang pinakaunang makikita ng organization representative pagkapasok sa portal. Dito pa lang, may quick overview na agad sila kung nasaan na sila sa buong compliance process.",
                "Makikita natin dito na hindi hiwa-hiwalay ang proseso. From profile setup, document submission, budget request, liquidation, hanggang YPOP incentive, nasa iisang portal lang ang progress tracking.",
                "Maganda itong ipakita sa panel dahil pinapakita nito na ang system ay hindi lang storage ng files, kundi guided workflow talaga para sa user.",
            ],
            key_points=[
                "Centralized ang user workflow sa iisang dashboard.",
                "Madaling makita ng user kung anong next step ang kailangan niyang gawin.",
                "Nababawasan ang confusion dahil may status cues na agad bago pa pumasok sa bawat module.",
            ],
        ),
        DemoSection(
            number="2",
            title="Organization Profile Setup and Verification",
            goal="Ipaliwanag na ang profile ang foundation ng buong compliance workflow at kailangan muna itong kumpleto bago makagalaw sa ibang modules.",
            actions=[
                "Buksan ang Organization Profile section.",
                "Ipakita ang required fields tulad ng organization name, email, contact number, district, barangay, classification, advocacy, adviser, representative, at address.",
                "Kung available sa demo data, ipakita rin ang difference ng existing organization versus bagong organization flow.",
                "I-save ang profile at ipaliwanag ang pending review or verified status.",
            ],
            script_lines=[
                "Bago payagan ang organization na mag-upload ng documents o mag-request ng budget, kailangan munang ma-complete ang organization profile.",
                "Mapapansin dito na hindi lang simpleng name at contact details ang hinihingi. Kasama rin ang classification, advocacy, at representative details dahil ito ang basehan ng admin para ma-validate ang legitimacy at category ng organization.",
                "Kapag na-save ang profile, hindi ibig sabihin automatic approved na agad. Dumadaan pa rin ito sa review process ng admin, kaya may malinaw na profile status gaya ng pending review, verified, o needs update.",
                "Mahalagang banggitin sa demo na itong step na ito ang nagse-set ng data quality para sa mga susunod na transactions sa system.",
            ],
            key_points=[
                "Profile completeness ang unang gate bago makapag-comply sa ibang modules.",
                "May admin validation layer, kaya controlled ang data quality.",
                "Clear ang distinction ng incomplete, pending review, verified, at needs update states.",
            ],
        ),
        DemoSection(
            number="3",
            title="Document Submissions and OCR-Assisted Review",
            goal="Ipakita kung paano sinusupport ng portal ang structured document submission at user-side validation bago ito makarating sa admin.",
            actions=[
                "Buksan ang Document Submissions section.",
                "Ipakita ang listahan ng required templates tulad ng Constitution and By-Laws, Form B, officers and adviser list, members in good standing, Form A, at data request form.",
                "Mag-open ng isang upload slot at i-explain na PDF ang default, habang may specific slots na puwedeng PDF or XLSX.",
                "Ipakita ang OCR review modal kung available sa demo data, pati ang confirmation bago mag-submit for admin review.",
                "Ipakita rin ang approved, needs revision, or rejected states kung merong sample entries.",
            ],
            script_lines=[
                "Sa document submission flow, hindi pinapasa ng system ang burden sa user nang walang guidance. Nakalista na agad ang required documents at may corresponding template slots pa para mas malinaw kung ano ang kulang.",
                "Isang importanteng feature dito ay ang OCR-assisted review. Bago tuluyang i-submit ang file, may pagkakataon ang user na makita ang extracted text or detected details para ma-double check kung tama ang na-upload.",
                "Sa practical use, malaking tulong ito para mabawasan ang mali, maling file, o unreadable submissions bago pa ito i-review ng admin.",
                "Kapag naisumite na, ang document ay magkakaroon ng admin-facing status tulad ng under review, needs revision, approved, o rejected. Ibig sabihin, transparent sa user kung ano ang nangyayari sa submission niya.",
            ],
            key_points=[
                "May guided list ng documentary requirements.",
                "May OCR-assisted checking bago final submission.",
                "Transparent ang review lifecycle sa user through visible submission statuses and remarks.",
            ],
        ),
        DemoSection(
            number="4",
            title="Budget Request Workflow",
            goal="Ipaliwanag ang controlled request process para sa funding at kung paano ito naka-link sa profile at supporting documents.",
            actions=[
                "Buksan ang Budget Requests section.",
                "Ipakita ang form fields tulad ng activity title, description, date, venue, requested amount, purpose category, at remarks.",
                "Ipakita na required ang attached budget file bago makumpleto ang submission.",
                "Kung may sample data, i-demo ang status progression from draft to submitted, under review, approved for FTF, hard copy submitted, at budget released.",
                "Ipakita rin kung paano nababasa ng user ang admin remarks kapag needs revision o rejected ang request.",
            ],
            script_lines=[
                "Pagdating naman sa budget request, ang portal ay hindi lang simpleng request form. Naka-structure ito para masigurong may malinaw na activity details, financial amount, at supporting document bago umabot sa admin evaluation.",
                "Dito natin makikita na may lock conditions din ang system. Halimbawa, kapag approved na for face-to-face submission or budget release, hindi na puwedeng basta baguhin ng user ang record.",
                "Sa side ng user, mahalagang makita ang transparency ng budget request history. Kapag may revision na hinihingi ang admin, kitang-kita rin ang remarks para alam ng organization kung ano ang dapat i-correct.",
                "Ang strength ng flow na ito ay auditability. Hindi nawawala sa user ang context ng request, at hindi rin umaasa sa informal chat or manual reminder lang ang process.",
            ],
            key_points=[
                "Budget requests are structured and attachment-driven.",
                "May visible status history at admin remarks for revisions.",
                "May control gates para maiwasan ang unauthorized edits after approval milestones.",
            ],
        ),
        DemoSection(
            number="5",
            title="Liquidation Reporting",
            goal="Ipakita ang post-activity accountability flow pagkatapos ma-release ang budget.",
            actions=[
                "Buksan ang Liquidation Reports section.",
                "I-explain na nagiging available lang ang liquidation kapag eligible na ang linked budget request, lalo na pagkatapos ng release status.",
                "Ipakita ang report form, attachments, at timeline-related statuses gaya ng submitted, under review, needs revision, hard copy submitted, completed liquidated, o overdue.",
                "Kung may sample record, buksan ang notes or remarks para maipakita ang admin feedback loop.",
            ],
            script_lines=[
                "Matapos ma-approve at ma-release ang budget, hindi doon nagtatapos ang workflow. May liquidation reporting module para sa post-activity accountability.",
                "Mahalagang i-highlight na conditional ang access dito. Hindi lahat ng user puwedeng gumawa agad ng liquidation; available lang ito kapag ang budget request ay nasa tamang stage na.",
                "Sa actual process, dito isi-submit ng organization ang liquidation details at files, at makikita rin nila kung tinanggap na ba ito, kailangan bang i-revise, o overdue na ang submission.",
                "Sa panel presentation, magandang idiin na ang portal ay end-to-end. Hindi lang niya pinapadali ang paghingi ng pondo, kundi pati ang pag-close ng financial compliance cycle.",
            ],
            key_points=[
                "Naka-link ang liquidation sa actual budget lifecycle.",
                "May accountability trail after activity implementation.",
                "Conditional ang access at malinaw ang completion or overdue states.",
            ],
        ),
        DemoSection(
            number="6",
            title="YPOP Incentive Submission",
            goal="Ipaliwanag ang semester-based incentive workflow at kung paano sinusukat ng system ang qualification.",
            actions=[
                "Buksan ang YPOP Incentive section.",
                "Ipakita ang current or past semester entries, including status, points earned, points required, at validation deadline.",
                "Kung may sample draft or qualified entry, ipakita ang city-led attendance, org-led project count, attached proof files, at admin remarks.",
                "I-explain na may relationship ang YPOP flow sa budget request kapag incentive-linked ang PPA.",
            ],
            script_lines=[
                "Ang YPOP Incentive module ay para sa organizations na gustong ma-track ang qualification nila based on participation and organization-led activities.",
                "Makikita rito na hindi lang basta upload ng proof ang system. Naka-structure ang scoring logic batay sa city-led attendance, organization-led projects, at required points per semester.",
                "Kapag na-validate ng admin, malinaw kung qualified o not qualified ang organization, at may remarks din kung kailangan ng revision o dagdag na ebidensya.",
                "Magandang banggitin din na sa current workflow, may connection ang YPOP record sa budget request kapag incentive-based ang request type. Ibig sabihin, integrated ang grants and compliance logic ng portal.",
            ],
            key_points=[
                "Semester-based at criteria-driven ang YPOP validation flow.",
                "Readable sa user ang points, deadlines, status, at admin remarks.",
                "Integrated ang incentive workflow sa broader funding and compliance process.",
            ],
        ),
        DemoSection(
            number="7",
            title="News Releases and Updates",
            goal="Ipakita na may information channel din ang portal para sa latest updates at official announcements.",
            actions=[
                "Buksan ang News Releases section.",
                "Ipakita ang listahan ng published announcements at kung paano ito nababasa ng user sa portal.",
                "I-mention na ang content na ito ay admin-managed kaya controlled ang official communication.",
            ],
            script_lines=[
                "Bukod sa transactional modules, may News Releases section din ang user portal para hindi hiwa-hiwalay ang communication ng LYDO sa organizations.",
                "Sa halip na umasa sa external channels lang, may official update area mismo sa system kung saan makikita ng users ang latest announcements, reminders, at release information.",
                "Mahalaga ito sa user experience dahil pinagsasama nito ang action workflow at information updates sa isang platform.",
            ],
            key_points=[
                "Built-in official communication channel ang portal.",
                "Admin-managed ang updates kaya mas reliable at consistent ang announcements.",
                "Mas less fragmented ang user experience dahil hindi na kailangang lumabas pa ng system para sa core updates.",
            ],
        ),
        DemoSection(
            number="8",
            title="Notifications and User Awareness",
            goal="Tapusin ang walkthrough sa pagpapakita ng reminders at status awareness features na tumutulong sa user na hindi makaligtaan ang important actions.",
            actions=[
                "Buksan ang Notifications section o notification indicators kung visible sa current layout.",
                "Ipakita ang examples ng alerts gaya ng profile reminder, document review status, released budget, o overdue liquidation.",
                "I-demo ang read and unread behavior kung available sa sample data.",
            ],
            script_lines=[
                "Ang notifications ang nagsisilbing follow-through mechanism ng portal. Kahit hindi araw-araw magbukas ang user ng bawat module, may reminders pa rin siya sa important updates at pending actions.",
                "Makakatulong ito lalo na sa compliance-heavy workflows dahil napapaalalahanan ang organization kung may pending review, released budget, o overdue liquidation na kailangang aksyunan.",
                "Sa demo, magandang i-emphasize na ang system ay proactive, hindi lang reactive. Hindi lang ito naghihintay na hanapin ng user ang status; ipinapakita rin nito ang status changes sa tamang oras.",
            ],
            key_points=[
                "Notifications improve follow-through and reduce missed steps.",
                "Useful ito sa multi-stage workflows tulad ng documents, budget, at liquidation.",
                "Pinapalakas nito ang usability ng portal through timely reminders.",
            ],
        ),
    ]

    for section in sections:
        add_demo_section(document, section)

    add_closing_section(
        document,
        [
            "Dito po nagtatapos ang user side demo ng LYDO Connect.",
            "Sa kabuuan, makikita natin na ang user portal ay hindi lang simpleng online form repository. Isa itong guided system na tumutulong sa organizations mula profile setup, document compliance, funding request, liquidation, hanggang incentive qualification at official updates.",
            "Ang pinakamahalagang value ng user side ay clarity. Bawat step ay may malinaw na status, may visible feedback, at may structured flow kaya mas madali para sa user na maintindihan kung ano ang susunod nilang gagawin.",
        ],
    )

    return document


def build_admin_document() -> Document:
    document = Document()
    set_page_layout(document)
    set_styles(document)

    add_title_block(
        document,
        "LYDO Connect Admin Side Demo Script",
        "Comprehensive presenter script for the administrative control portal",
        "Student presenters, panel defense speakers, admin onboarding leads, and documentation teams",
        "To explain how admins review registrations, validate submissions, manage content, monitor finance workflows, and maintain accountability logs.",
    )

    add_note_box(
        document,
        "How to Use This Script",
        [
            "Ang admin side demo ay best na i-present bilang control center ng system, hindi lang bilang listahan ng admin pages.",
            "Ang current admin modules na sakop ng script na ito ay Overview, Registration Review, YORP Registry, Budget Requests, Liquidation Reports, Budget Tracking, YPOP Validation, News Releases, Templates, Notifications, at Activity Logs.",
            "Kapag may module na walang sample transaction sa demo data, gamitin ang expected admin responsibility lines para malinaw pa rin ang explanation.",
        ],
    )

    add_coverage_table(
        document,
        [
            ("1", "Overview", "Portal-wide visibility and quick status scanning"),
            ("2", "Registration Review", "Validation of profiles and documentary submissions"),
            ("3", "YORP Registry", "Registry maintenance and organization records"),
            ("4", "Budget Requests", "Funding review and approval decision points"),
            ("5", "Liquidation Reports", "Post-release review and financial accountability"),
            ("6", "Budget Tracking", "Monitoring utilization and barangay allocation"),
            ("7", "YPOP Validation", "Semester incentive evaluation and scoring review"),
            ("8", "News Releases", "Official content publishing for users"),
            ("9", "Templates", "Managing required compliance templates"),
            ("10", "Notifications and Activity Logs", "System awareness and traceability"),
        ],
    )

    sections = [
        DemoSection(
            number="1",
            title="Overview Dashboard",
            goal="Ipakita ang admin portal bilang command center kung saan agad nakikita ang high-level system status at urgent items.",
            actions=[
                "Mag-sign in bilang admin at buksan ang Overview page.",
                "I-tour ang sidebar groups: Organizations, Budget & Finance, Content, at System.",
                "I-highlight ang overview cards, pending counts, o quick summaries na available sa dashboard.",
            ],
            script_lines=[
                "Sa admin side, ang Overview ang nagsisilbing command center ng buong LYDO Connect system.",
                "Dito pa lang, may immediate visibility na ang admin sa registrations, budget workflows, liquidation follow-ups, content updates, at iba pang operational concerns.",
                "Magandang i-frame ito sa panel bilang evidence na ang admin portal ay ginawa para sa monitoring at decision support, hindi lang para sa manual encoding ng data.",
            ],
            key_points=[
                "Centralized ang monitoring view ng admin.",
                "Nakakatulong ang overview sa prioritization ng urgent workflows.",
                "Pinapakita nito na ang system ay operational tool, hindi static records page lang.",
            ],
        ),
        DemoSection(
            number="2",
            title="Registration Review",
            goal="Ipakita ang validation workflow ng admin para sa organization profile at documentary compliance.",
            actions=[
                "Buksan ang Registration Review module.",
                "Pumili ng isang organization record at i-open ang profile details.",
                "Ipakita ang submitted files, preview options, status controls, at admin remarks area.",
                "I-demo ang possible actions gaya ng verify profile, approve document, mark for revision, o reject kapag may issue.",
            ],
            script_lines=[
                "Dito sa Registration Review, nakikita ng admin ang incoming organization applications at compliance submissions.",
                "Hindi lang basta listahan ang module na ito. Makikita rin dito ang profile details, attached files, review remarks, at status controls para sa bawat organization.",
                "Kapag valid ang requirements, puwedeng i-verify ng admin ang profile at i-approve ang mga document submissions. Kapag may kulang o may problema, puwede rin itong ibalik as needs revision o i-reject with remarks.",
                "Sa perspective ng governance, mahalaga ito dahil dito nangyayari ang quality control bago maging fully active ang organization sa system.",
            ],
            key_points=[
                "May structured review process para sa profiles at documents.",
                "Merong visible remarks and status controls para sa admin decisions.",
                "Nagsisilbi itong quality-control gateway ng user onboarding and compliance flow.",
            ],
        ),
        DemoSection(
            number="3",
            title="YORP Registry Management",
            goal="Ipakita ang admin responsibility sa master registry ng organizations.",
            actions=[
                "Buksan ang YORP Registry module.",
                "Ipakita ang organization records at anumang searchable or filterable registry view na available.",
                "I-explain ang value nito bilang maintained source of truth ng accredited or tracked organizations.",
            ],
            script_lines=[
                "Ang YORP Registry ang mas formal na record-keeping side ng organizations sa system.",
                "Habang ang Registration Review ay mas workflow-oriented, ang registry ay tumutulong naman sa admin na mag-maintain ng organized list ng youth organizations at kanilang current standing.",
                "Mahalagang i-highlight na ito ang nagiging reliable reference kapag may reporting, validation, o follow-up na kailangang gawin across organizations.",
            ],
            key_points=[
                "Registry view supports master data management.",
                "Nakakatulong ito sa consistency ng organization records.",
                "Useful ito for validation, reporting, and administrative follow-up.",
            ],
        ),
        DemoSection(
            number="4",
            title="Budget Request Review and Approval",
            goal="Ipakita ang decision-making workflow ng admin sa pag-review ng funding requests.",
            actions=[
                "Buksan ang Budget Requests module sa admin side.",
                "Pumili ng sample request at ipakita ang activity details, requested amount, attached files, at review history kung available.",
                "I-demo ang possible admin actions gaya ng approve, request revision, reject, mark hard copy submitted, o mark budget released.",
                "Ipakita rin ang remarks area bilang formal feedback channel sa user.",
            ],
            script_lines=[
                "Sa admin side ng budget workflow, dito nangyayari ang actual evaluation ng funding requests na galing sa organizations.",
                "Makikita ng admin ang detalye ng activity, amount requested, attached supporting document, at current status ng bawat request.",
                "Ang mahalagang point dito ay may controlled progression ang status. Hindi random ang galaw ng request; dumadaan ito sa review, approval milestone, hard copy stage, at release stage depende sa actual process.",
                "Kapag may issue, may remarks mechanism din para malinaw sa user kung bakit kailangan ng revision o bakit na-reject ang request.",
            ],
            key_points=[
                "Structured ang admin approval path ng budget requests.",
                "Each status change can carry formal remarks for accountability.",
                "Nire-reinforce nito ang controlled financial workflow sa system.",
            ],
        ),
        DemoSection(
            number="5",
            title="Liquidation Report Review",
            goal="Ipakita kung paano bina-balance ng admin ang support at accountability pagkatapos ng budget release.",
            actions=[
                "Buksan ang Liquidation Reports module.",
                "Mag-open ng sample liquidation report at ipakita ang linked budget request context.",
                "I-review ang report status, supporting attachments, at admin remark actions.",
                "I-mention ang overdue handling at completion states.",
            ],
            script_lines=[
                "Pagdating naman sa liquidation, ang role ng admin ay siguraduhin na ang na-release na pondo ay may maayos na post-activity documentation at liquidation evidence.",
                "Sa module na ito, makikita ang relation ng liquidation report sa original budget request, kaya buo ang financial trail ng activity.",
                "Puwedeng i-approve, ipa-revise, o i-mark as overdue ang report depende sa completeness at timeliness ng submission.",
                "Magandang ipakita sa demo na hindi natatapos ang admin control sa pag-release ng pondo; sinasaklaw din nito ang liquidation closure at accountability follow-through.",
            ],
            key_points=[
                "Linked ang liquidation review sa original funding request.",
                "May timeline-sensitive states tulad ng overdue and completed liquidated.",
                "Pinapalakas nito ang accountability and closure mechanism ng portal.",
            ],
        ),
        DemoSection(
            number="6",
            title="Budget Tracking and Barangay Allocation Monitoring",
            goal="Ipakita ang analytics and oversight layer ng admin sa budget utilization across requests and locations.",
            actions=[
                "Buksan ang Budget Tracking module.",
                "I-tour ang overview view at barangay allocation tab kung available.",
                "Ipakita ang totals, released amounts, remaining balances, utilization indicators, at filters by district or barangay.",
                "I-mention ang export capability kung visible sa current build.",
            ],
            script_lines=[
                "Bukod sa individual budget approvals, may separate monitoring view din ang admin para makita ang bigger picture ng budget utilization.",
                "Dito makikita kung alin ang on track, alin ang nangangailangan ng attention, at alin ang may overdue or incomplete liquidation behavior.",
                "May allocation perspective rin per barangay, kaya hindi lang activity-level ang analysis kundi location-based din kung saan napupunta ang pondo.",
                "Sa panel, magandang i-emphasize na ito ang decision-support layer ng admin portal dahil nagbibigay ito ng summarized and exportable oversight data.",
            ],
            key_points=[
                "May admin analytics view for utilization and allocation.",
                "Useful ito for oversight beyond individual transactions.",
                "Supports exportable and filterable monitoring for reporting needs.",
            ],
        ),
        DemoSection(
            number="7",
            title="YPOP Validation",
            goal="Ipakita ang admin-side evaluation ng incentive submissions and semester configuration.",
            actions=[
                "Buksan ang YPOP Validation module.",
                "Kung may semester list, ipakita ang creation or selection ng semester period.",
                "I-demo ang review ng organization submissions, attendance records, org-led project count, points, at final status.",
                "Ipakita ang admin remarks at qualified or not qualified decision state.",
            ],
            script_lines=[
                "Sa YPOP Validation, dito tinitingnan ng admin kung pasado ba ang organization sa requirements ng incentive program.",
                "Hindi lang manual judgment ang gamit dito. Naka-anchor ang review sa semester setup, activity points, attendance data, at org-led project count.",
                "Makikita rin ng admin kung qualified, under review, needs revision, o not qualified ang submission, kaya transparent at criteria-based ang decision.",
                "Kung gusto ninyong idiin ang value ng system, magandang sabihin na ang module na ito ay nagbabawas ng ambiguity sa incentive evaluation process.",
            ],
            key_points=[
                "Criteria-based ang validation ng YPOP submissions.",
                "May admin control over semester setup and review outcomes.",
                "Nakakatulong ito para standardized ang incentive qualification decisions.",
            ],
        ),
        DemoSection(
            number="8",
            title="News Releases Management",
            goal="Ipakita kung paano kino-control ng admin ang official announcements na nakikita ng user portal.",
            actions=[
                "Buksan ang News Releases module.",
                "Ipakita ang create or edit form para sa title, description, post date, link, at visibility status.",
                "I-demo ang draft, published, at hidden states kung available.",
            ],
            script_lines=[
                "Sa content side, may News Releases management ang admin para sa official announcements ng LYDO Connect.",
                "Dito puwedeng gumawa, mag-edit, mag-publish, o mag-hide ng updates na makikita sa user portal.",
                "Ang magandang takeaway dito ay integrated ang communication management sa same admin environment, kaya hindi hiwalay ang operational workflow at official announcements.",
            ],
            key_points=[
                "Admin-controlled ang publication workflow ng updates.",
                "May draft and visibility states for content governance.",
                "Naka-integrate ang communications sa rest ng administrative portal.",
            ],
        ),
        DemoSection(
            number="9",
            title="Templates Management",
            goal="Ipakita kung paano mina-manage ng admin ang required compliance templates na ginagamit ng users.",
            actions=[
                "Buksan ang Templates module.",
                "Ipakita ang listahan ng required document types at kung may template file upload or replacement options.",
                "I-explain na dito nanggagaling ang guidance structure ng user-side document submission flow.",
            ],
            script_lines=[
                "Ang Templates module ang nagbibigay ng backbone sa document submission process ng user side.",
                "Kapag may bagong template file o updated requirement, dito ito ina-upload at mina-manage ng admin para manatiling consistent ang compliance package na sinusunod ng users.",
                "Mahalagang ipakita ito dahil pinapakita nito na configurable ang system at hindi hard-coded lang ang guidance sa document requirements.",
            ],
            key_points=[
                "Admin can manage the required template ecosystem.",
                "Supports consistency of documentary compliance requirements.",
                "Improves maintainability when requirements change over time.",
            ],
        ),
        DemoSection(
            number="10",
            title="Notifications and Activity Logs",
            goal="Tapusin ang admin walkthrough sa accountability and traceability features ng system.",
            actions=[
                "Buksan ang Notifications section at ipakita ang unread or recent items.",
                "Buksan ang Activity Logs at ipakita ang recorded system actions kung available.",
                "I-mention na ang logs ay useful para malaman kung sino ang gumawa ng action at kailan ito nangyari.",
            ],
            script_lines=[
                "Sa huling bahagi ng admin demo, mahalagang ipakita ang notifications at activity logs dahil dito makikita ang operational awareness at accountability trail ng system.",
                "Ang notifications ay tumutulong sa admin na makita ang pending actions, gaya ng bagong submissions o overdue reports.",
                "Samantala, ang activity logs naman ang nagbibigay ng traceability kung anong admin action ang nangyari, kailan ito ginawa, at anong record ang naapektuhan.",
                "Para sa panel, magandang idiin na ang dalawang features na ito ang nagpapalakas sa governance aspect ng LYDO Connect.",
            ],
            key_points=[
                "Notifications support timely admin action.",
                "Activity logs strengthen auditability and accountability.",
                "These features make the portal more reliable for real operations.",
            ],
        ),
    ]

    for section in sections:
        add_demo_section(document, section)

    add_closing_section(
        document,
        [
            "Dito po nagtatapos ang admin side demo ng LYDO Connect.",
            "Sa kabuuan, ang admin portal ang nagsisilbing operational and governance backbone ng system. Dito mina-manage ang registrations, validations, budget workflows, liquidation, incentives, content updates, at accountability logs.",
            "Ang pinakamahalagang mensahe ng admin side ay control with traceability. Hindi lang nito pinapabilis ang trabaho ng admin, kundi tinitiyak din nito na ang bawat mahalagang action ay may malinaw na process, status, at audit trail.",
        ],
    )

    return document


def save_document(document: Document, output_name: str) -> Path:
    path = Path(output_name)
    document.save(path)
    return path.resolve()


def main() -> None:
    outputs = [
        save_document(build_user_document(), "LYDO_Connect_User_Side_Comprehensive_Script_Taglish.docx"),
        save_document(build_admin_document(), "LYDO_Connect_Admin_Side_Comprehensive_Script_Taglish.docx"),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
